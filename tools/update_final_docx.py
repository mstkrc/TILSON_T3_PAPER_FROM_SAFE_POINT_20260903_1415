from pathlib import Path
from docx import Document

root = Path(__file__).parents[1]
docs = sorted((root / "DOKUMANTASYON").glob("*.docx"))
addendum = [
    "FINAL HANDOFF / DOCUMENTATION CLOSURE",
    "Faz-0 → Faz-20: PASS / LOCKED.",
    "KONU-1 → KONU-49: LOCKED. KONU-49: openpyxl 3.1.5 yalnız Faz-14 Report/Excel Export için onaylıdır.",
    "Faz-19 full regression: 77/77 PASS.",
    "Live: LIVE_TRADING=false; aktif LIVE_TRADING=true yok; gerçek emir ve Binance order endpoint yok.",
    "Paper-only korunur; Ledger single source of truth olarak kalır.",
    "Faz-20 final handoff/documentation closure tamamlandı.",
    "Final Word/DOC paketinin bu güncel kapanış kayıtlarıyla toplu güncellenmesi sonraki kullanıcı onaylı adımdır.",
]
for path in docs:
    document = Document(path)
    document.add_page_break()
    for index, text in enumerate(addendum):
        document.add_heading(text, level=1 if index == 0 else 2) if index == 0 else document.add_paragraph(text)
    document.save(path)
print(f"UPDATED={len(docs)}")
